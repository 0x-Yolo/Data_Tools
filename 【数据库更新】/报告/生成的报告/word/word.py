from docx import Document
from docx.shared import Inches, Cm, Pt,RGBColor
from fig_organize import merge
from docx.oxml.ns import qn
from primary_market_plot import GK
import ReportGenerator as rg

report = rg.weeklyReport()
base = '2021-06-04'
end = '2021-06-11'
start = '2021-06-07'
# * 流动性
r_dr = report.r_dr('2019-07-01')
mone = report.monetary_policy_tools(base, end)
# * 利率债
gk_gz = report.gk_gz('2015-01-01')
issue = report.prmy_mkt_weekly_issue(start, end)
gk = GK([5,10])
net_rate, net_credit = report.most_net_buy_amt(start, end)
term = report.term_spread()
# * 信用债
credit_figs = report.secondary_credit('2021-05-31', end)



doc = Document()
doc.styles['Normal'].font.name = 'STKaiti'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'STKaiti')
#changing the page margins
sections = doc.sections
margin = 1.27
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
# * para1
Head = doc.add_heading('', 1)
run  = Head.add_run("流动性")
run.font.name='STKaiti'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'STKaiti')
run.font.color.rgb = RGBColor(0,0,0)
# r-dr\净投放
doc.add_picture(merge([r_dr,mone]), width=Cm(18.46))
para_1 = doc.add_paragraph('流动性观察')
paragraph_format = para_1.paragraph_format
paragraph_format.first_line_indent = Inches(-0.25)

# * para2
Head = doc.add_heading('', 1)
run  = Head.add_run("利率债")
run.font.name='STKaiti'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'STKaiti')
run.font.color.rgb = RGBColor(0,0,0)
# gk-gz\一级发行
doc.add_picture(merge([gk_gz,issue]), width=Cm(18.46))
# gk综收\利率债净买入
doc.add_picture(merge([gk, net_rate]), width = Cm(18.46))
# 期限利差
doc.add_picture(merge(term), width = Cm(18.46))
para_2 = doc.add_paragraph('利率债交易情绪与配置力量')
paragraph_format = para_2.paragraph_format
paragraph_format.first_line_indent = Inches(-0.25)

# * para3
Head = doc.add_heading('', 1)
run  = Head.add_run("信用债")
run.font.name='STKaiti'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'STKaiti')
run.font.color.rgb = RGBColor(0,0,0)
# 信用债配置力量与风险偏好 
doc.add_picture(merge([credit_figs[1],credit_figs[2]]), width = Cm(18.46))
# 信用债净买入
doc.add_picture(merge([net_credit]), width = Cm(18.46 / 2))
para_3 = doc.add_paragraph('利率债交易情绪与配置力量')
paragraph_format = para_3.paragraph_format
paragraph_format.first_line_indent = Inches(-0.25)

doc.save('1.docx')

# 1.流动性
## P1 R-DR
## P2 超储率
## P3 净投放


# 2.利率债
## 国开-国债
## 配置力量
## 期限利差
## 现券净买入 2个期限

# 3.信用债
## 配置力量
## 风险偏好
## 净买入最活跃的机构





# doc = Document()
# doc.add_heading('Doc Title', 1)

# doc.add_picture('/Users/wdt/Desktop/tpy/Data_Tools/【数据库更新】/报告/生成的报告/周报_0530/MLF与同业存单.jpg',\
#     width=Inches(2.15))
# doc.add_picture('/Users/wdt/Desktop/tpy/Data_Tools/【数据库更新】/报告/生成的报告/周报_0530/MLF与同业存单.jpg',\
# width=Inches(2.15))
# doc.save('test.docx')