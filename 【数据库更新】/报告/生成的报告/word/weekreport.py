import data_organize as do
import ReportGenerator as rg
from docx import Document
from docx.shared import Inches, Cm, Pt,RGBColor
from fig_organize import merge
from docx.oxml.ns import qn
from primary_market_plot import GK


# * 周报 
weekreport = rg.weeklyReport(isMonth = False)
base = '2021-06-04'
end = '2021-06-11'
start = '2021-06-07'
## ! 含有base_day
cash = weekreport.cash_cost(base, end)
mone = weekreport.monetary_policy_tools(base, end)
dps = weekreport.interbank_deposit(base, end)
bp_change = weekreport.rates_change(base, end)#bp
## * 二级 :近两周 
second_credit = weekreport.secondary_credit('2021-05-31', end)
second_rate = weekreport.secondary_rate('2021-05-31', end)
## * 一级 :近两周 
prmy_issue = weekreport.prmy_mkt_weekly_issue(start, end)
## * 现券
net1 = weekreport.fig_net_data(start, end)
## * 综收 2020以来
prmy_senti_GK = GK([5,10])


doc = Document()
doc.styles['Normal'].font.name = 'STKaiti'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'STKaiti')
#changing the page margins
sections = doc.sections
margin = 1.27
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

doc.add_picture(merge([cash,mone]), width=Cm(18.46))
doc.add_picture(merge([dps,prmy_issue]), width=Cm(18.46))

doc.add_picture(merge([prmy_senti_GK]), width=Cm(18.46))

doc.add_picture(merge([bp_change[0],bp_change[1]]), width=Cm(18.46))
doc.add_picture(merge([bp_change[2],second_credit[0]]), width=Cm(18.46))
doc.add_picture(merge([bp_change[2],second_credit[0]]), width=Cm(18.46))
doc.add_picture(merge([net1[0],net1[1]]), width=Cm(18.46))

doc.save('1.docx')